"""
DescomplicaCV - API Backend em Python com FastAPI

Esta API fornece endpoints para:
1. Receber arquivos de currículo em diferentes formatos (PDF, DOCX, TXT)
2. Processar estes arquivos e extrair informações relevantes
3. Gerar um novo currículo em formato PDF padronizado

Desenvolvido como parte do projeto DescomplicaCV.

Requisitos:
- Python 3.7+
- FastAPI
- Uvicorn (servidor ASGI)
- PyMuPDF (manipulação de PDFs)
- python-docx (manipulação de DOCXs)

Instalação:
pip install fastapi uvicorn python-multipart pymupdf python-docx
"""

import os
import io
import json
from fastapi import FastAPI, UploadFile, HTTPException, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse
import fitz  # PyMuPDF - Biblioteca para manipulação de PDFs
from docx import Document  # python-docx - Biblioteca para manipulação de arquivos DOCX
import spacy  # Biblioteca para processamento de linguagem natural
from spacy.matcher import Matcher  # Matcher para busca de padrões no texto


# Criando a instância do FastAPI com metadados
app = FastAPI(
    title="API DescomplicaCV", 
    version="0.1.0",
    description="API para conversão de currículos para o formato PDF padronizado"
)

# Carregando o modelo spaCy para português
try:
    nlp = spacy.load("pt_core_news_sm")
    print("Modelo spaCy carregado com sucesso!")
except Exception as e:
    print(f"Erro ao carregar o modelo spaCy: {str(e)}")
    nlp = None  # Se houver erro, define como None para tratamento posterior

# Inicializando o matcher do spaCy para identificar padrões relevantes
matcher = Matcher(nlp.vocab) if nlp else None

# Adicionar padrões para encontrar informações relevantes em currículos
if matcher:
    # Padrão para e-mail
    matcher.add("EMAIL", [[{"LIKE_EMAIL": True}]])
    
    # Padrão para telefone (simplificado)
    matcher.add("TELEFONE", [
        [{"SHAPE": "dd-dddd-dddd"}],
        [{"SHAPE": "ddddd-dddd"}],
        [{"SHAPE": "(dd)dddd-dddd"}],
        [{"SHAPE": "(dd)ddddd-dddd"}],
        [{"TEXT": {"REGEX": r"\(\d{2}\)\s?\d{4,5}-?\d{4}"}}]
    ])
    
    # Padrão para educação
    matcher.add("EDUCACAO", [
        [{"LOWER": {"IN": ["graduação", "graduacao", "formação", "formacao", "bacharel", "bacharelado", "licenciatura"]}}, 
         {"IS_ALPHA": True, "OP": "*"}],
        [{"LOWER": {"IN": ["mestrado", "doutorado", "especialização", "especializacao", "pós-graduação", "pos-graduacao"]}}, 
         {"IS_ALPHA": True, "OP": "*"}]
    ])
    
    # Padrão para experiência
    matcher.add("EXPERIENCIA", [
        [{"LOWER": {"IN": ["experiência", "experiencia", "profissional", "trabalhou", "atuou"]}}, 
         {"IS_ALPHA": True, "OP": "*"}]
    ])

# Configuração do CORS (Cross-Origin Resource Sharing)
# Permite que o frontend Vue.js (rodando em outra porta) acesse esta API
origins = [
    "http://localhost:5173",  # Endereço comum do Vue com Vite em desenvolvimento
    "http://localhost:3000",  # Porta alternativa comum para desenvolvimento
    # Adicione aqui o endereço do seu frontend em produção quando souber
]

# Adicionando o middleware CORS à aplicação
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],  # Permite todos os métodos HTTP (GET, POST, etc.)
    allow_headers=["*"],  # Permite todos os cabeçalhos HTTP
)
#=========================================================================================

@app.post("/upload_cv")
async def upload_cv(file: UploadFile = File(...)):
    """
    Endpoint para receber e processar arquivos de currículo em diferentes formatos.
    
    Aceita arquivos PDF, DOCX e TXT, processando-os de acordo com o tipo.
    Utiliza spaCy para analisar o conteúdo e extrair informações relevantes.
    
    Parameters:
        file (UploadFile): O arquivo de currículo enviado pelo cliente
        
    Returns:
        dict: Informações extraídas do currículo e metadados do arquivo
        
    Raises:
        HTTPException: Em caso de erro no processamento do arquivo
    """
    filename = file.filename
    content = await file.read()
    
    # Determina a extensão do arquivo
    ext = os.path.splitext(filename)[1].lower()
    
    # Processa baseado no tipo de arquivo
    if ext == ".pdf":
        try:
            # Abre o PDF usando PyMuPDF (fitz) para validação
            # stream=content carrega o PDF a partir do conteúdo em memória
            pdf = fitz.open(stream=content, filetype="pdf")
            num_pages = len(pdf)
            
            # Extrair o texto completo de todas as páginas do PDF
            texto_completo = ""
            for page_num in range(num_pages):
                page = pdf[page_num]
                texto_completo += page.get_text()
            
            # Informações adicionais do documento
            metadata = pdf.metadata
            
            # Extrair informações usando spaCy
            info_extraidas = extrair_informacoes_cv(texto_completo)
            
            return {
                "filename": filename,
                "format": "pdf",
                "pages": num_pages,
                "title": metadata.get("title", "Sem título"),
                "author": metadata.get("author", "Autor desconhecido"),
                "texto_completo": texto_completo[:500] + "..." if len(texto_completo) > 500 else texto_completo,
                "info_extraidas": info_extraidas,
                "message": "PDF analisado com sucesso! Informações extraídas."
            }
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Arquivo PDF inválido ou corrompido: {str(e)}")
    elif ext == ".docx":
        try:
            # Cria um objeto BytesIO para trabalhar com o arquivo em memória
            docx_bytes = io.BytesIO(content)
            
            # Carrega o documento DOCX na memória
            doc = Document(docx_bytes)
            
            # Extrai o texto completo de todos os parágrafos
            texto_completo = "\n".join([p.text for p in doc.paragraphs])
            
            # Informações adicionais do documento
            paragraphs_count = len(doc.paragraphs)
            
            # Extrair informações usando spaCy
            info_extraidas = extrair_informacoes_cv(texto_completo)
            
            return {
                "filename": filename,
                "format": "docx",
                "paragraphs": paragraphs_count,
                "texto_completo": texto_completo[:500] + "..." if len(texto_completo) > 500 else texto_completo,
                "info_extraidas": info_extraidas,
                "message": "DOCX analisado com sucesso! Informações extraídas."
            }
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Erro ao processar DOCX: {str(e)}")
    elif ext == ".txt":
        try:
            # Converte o conteúdo binário para texto com codificação UTF-8
            texto_completo = content.decode("utf-8")
            lines = texto_completo.count("\n") + 1
            
            # Extrair informações usando spaCy
            info_extraidas = extrair_informacoes_cv(texto_completo)
            
            return {
                "filename": filename,
                "format": "txt",
                "lines": lines,
                "texto_completo": texto_completo[:500] + "..." if len(texto_completo) > 500 else texto_completo,
                "info_extraidas": info_extraidas,
                "message": "Arquivo TXT analisado com sucesso! Informações extraídas."
            }
            
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Erro ao processar TXT: {str(e)}")
    
    else:
        # Se o formato não for suportado, retorna erro 415 (Unsupported Media Type)
        raise HTTPException(
            status_code=415, 
            detail=f"Tipo de arquivo {ext} não suportado. Por favor, envie um arquivo PDF, DOCX ou TXT."
        )

def extrair_informacoes_cv(texto):
    """
    Extrai informações relevantes do texto do currículo usando spaCy.
    
    Args:
        texto (str): O texto completo do currículo
        
    Returns:
        dict: Dicionário contendo as informações extraídas
    """
    # Verifica se o spaCy foi inicializado corretamente
    if nlp is None:
        return {
            "erro": "Modelo spaCy não está disponível",
            "contatos": [],
            "educacao": [],
            "experiencia": [],
            "habilidades": []
        }
    
    # Processa o texto com spaCy
    doc = nlp(texto)
    
    # Dicionário para armazenar os resultados
    resultados = {
        "contatos": [],
        "educacao": [],
        "experiencia": [],
        "habilidades": []
    }
    
    # Extrair entidades nomeadas (pessoas, organizações, locais)
    for ent in doc.ents:
        if ent.label_ == "PER" and len(resultados.get("nome", "")) == 0:
            resultados["nome"] = ent.text
        elif ent.label_ == "ORG":
            if "organizacoes" not in resultados:
                resultados["organizacoes"] = []
            resultados["organizacoes"].append(ent.text)
        elif ent.label_ == "LOC":
            if "localizacoes" not in resultados:
                resultados["localizacoes"] = []
            resultados["localizacoes"].append(ent.text)
    
    # Usar o matcher para encontrar padrões específicos
    matches = matcher(doc)
    for match_id, start, end in matches:
        string_id = nlp.vocab.strings[match_id]  # Obtém o nome do padrão
        span = doc[start:end]  # O texto que correspondeu ao padrão
        
        if string_id == "EMAIL":
            resultados["contatos"].append({"tipo": "email", "valor": span.text})
        elif string_id == "TELEFONE":
            resultados["contatos"].append({"tipo": "telefone", "valor": span.text})
        elif string_id == "EDUCACAO":
            # Capturar o contexto ao redor da educação (3 tokens antes e depois)
            contexto_inicio = max(0, start - 3)
            contexto_fim = min(len(doc), end + 3)
            contexto = doc[contexto_inicio:contexto_fim].text
            resultados["educacao"].append(contexto)
        elif string_id == "EXPERIENCIA":
            # Capturar o contexto ao redor da experiência (3 tokens antes e depois)
            contexto_inicio = max(0, start - 3)
            contexto_fim = min(len(doc), end + 3)
            contexto = doc[contexto_inicio:contexto_fim].text
            resultados["experiencia"].append(contexto)
    
    # Identificar possíveis habilidades (substantivos e frases nominais não capturados em outras categorias)
    for chunk in doc.noun_chunks:
        # Se a chunk tiver entre 2 e 5 tokens, pode ser uma habilidade
        if 2 <= len(chunk) <= 5 and chunk.text.lower() not in [item.lower() for sublist in resultados.values() if isinstance(sublist, list) for item in sublist]:
            resultados["habilidades"].append(chunk.text)
    
    # Limitar o número de habilidades para evitar ruído
    if len(resultados["habilidades"]) > 10:
        resultados["habilidades"] = resultados["habilidades"][:10]
    
    # Remover duplicatas
    for key in resultados:
        if isinstance(resultados[key], list):
            resultados[key] = list(dict.fromkeys(resultados[key]))
    
    return resultados

#=========================================================================================
